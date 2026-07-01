from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page setup: A4, 2.5cm margins
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Helper to set font
def set_run(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # Force Times New Roman for East Asian fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_centered(text, size=12, bold=False, italic=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p

def add_body(space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    pf.first_line_indent = Cm(1.0)
    return p

# ===== TITLE =====
add_centered(
    "AI-Assisted Mobile Application for Visually Impaired Users:\nReal-Time Navigation, Object Retrieval, and Face Recognition",
    size=12, bold=True, space_before=0, space_after=6
)

# Author
add_centered("Cristian Gabriel FLOREA", size=12, bold=False, space_after=6)

# Affiliation
add_centered("Military Technical Academy \"Ferdinand I\", Bucharest, Romania",
             size=10, italic=True, space_after=2)

# Contact
add_centered("Contact: <CONTACT-EMAIL>", size=10, space_after=6)

# ===== INTRODUCTION =====
p = add_body(space_before=6, space_after=0)
r = p.add_run("Introduction")
set_run(r, size=12, bold=True)
r = p.add_run(": According to the World Health Organization, over 285 million people worldwide suffer from visual impairments, with approximately 39 million being completely blind. For these individuals, daily tasks such as independent navigation, locating personal objects, identifying people, and handling money represent significant challenges. Existing assistive applications like Seeing AI or Be My Eyes provide basic scene description and text recognition, but lack integrated spatial awareness, personalized object retrieval, and multimodal feedback tailored for fully autonomous use. This paper presents VisionAId, an Android application that combines multiple on-device AI models into a unified assistive system providing real-time navigation, personal object search with AR-guided navigation, face recognition, banknote identification, and color detection \u2014 all accessible through Romanian voice commands and multimodal feedback (text-to-speech, 3D spatial audio, and variable haptic feedback).")
set_run(r, size=12)

# ===== METHODS AND RESULTS =====
p = add_body(space_before=0, space_after=0)
r = p.add_run("Methods and Results")
set_run(r, size=12, bold=True)
r = p.add_run(": The application is built on a multi-layer ML architecture running entirely on-device using ONNX Runtime and LiteRT (TensorFlow Lite), ensuring privacy and offline operation. The ")
set_run(r, size=12)
r = p.add_run("navigation module")
set_run(r, size=12, italic=True)
r = p.add_run(" employs Depth Anything V2 Metric Small (ONNX, INT8 quantized) to produce metric depth maps at 378\u00d7378 resolution. A central corridor region (25% width \u00d7 45% height) of the depth map is continuously sampled for the minimum distance value, which triggers graduated haptic feedback through three alert levels: danger (<0.5\u2009m, strong vibration every 300\u2009ms), warning (0.5\u20131.0\u2009m, medium vibration every 500\u2009ms), and caution (1.0\u20131.5\u2009m, light vibration every 800\u2009ms). The ")
set_run(r, size=12)
r = p.add_run("object search module")
set_run(r, size=12, italic=True)
r = p.add_run(" integrates YOLOv11-Seg for real-time detection, MobileCLIP2-S2 for semantic embeddings, and ARCore for 3D anchor placement, enabling users to register personal objects through a guided multi-scale process and later retrieve them via voice commands with step-by-step spatial audio navigation. The ")
set_run(r, size=12)
r = p.add_run("face recognition module")
set_run(r, size=12, italic=True)
r = p.add_run(" uses YuNet (233\u2009KB ONNX model) for face detection and MobileFaceNet for 512-dimensional embedding extraction, implementing a FaceID-style registration with head rotation tracking across 8 angular buckets to capture up to 15 diverse embeddings per person. Additional modules include a custom-trained YOLOv12n model for Romanian banknote detection and a Gemini Vision API integration for contextual scene description in Romanian.")
set_run(r, size=12)

# Results paragraph
p = add_body(space_before=0, space_after=0)
r = p.add_run("Testing on a Samsung Galaxy S21 Ultra 5G demonstrates real-time performance: depth estimation runs at 3\u20135 FPS with proximity alerts, object detection achieves ~27 FPS, and face detection operates at ~15 FPS. The object retrieval pipeline achieves 80%+ matching accuracy using cluster-based embedding comparison with temporal consistency checking, while face recognition reaches a cosine similarity threshold of 0.45 for reliable identification. All six ML models (Depth Anything V2, YOLOv11-Seg, MobileCLIP2-S2, YuNet, MobileFaceNet, YOLOv12n) run concurrently on mobile hardware without GPU dependency, leveraging multi-threaded CPU inference with optimized memory management.")
set_run(r, size=12)

# ===== CONCLUSIONS =====
p = add_body(space_before=0, space_after=0)
r = p.add_run("Conclusions")
set_run(r, size=12, bold=True)
r = p.add_run(": VisionAId demonstrates that a comprehensive AI-assisted system for visually impaired users can run entirely on a standard Android smartphone, eliminating the need for cloud processing, specialized hardware, or external accessories (Bluetooth tags, RFID labels). The multi-layer architecture \u2014 combining depth estimation, object detection, semantic matching, face recognition, and augmented reality \u2014 provides a unified assistive experience through a fully voice-controlled Romanian interface. The depth-based navigation module offers continuous obstacle awareness without requiring object classification, while the personalized object search enables users to register and locate their own belongings through natural voice interaction. Future work will focus on extending the system with indoor SLAM-based mapping for persistent spatial memory, integrating on-device large language models for more natural conversational interaction, and conducting usability studies with visually impaired users to validate real-world effectiveness.")
set_run(r, size=12)

# ===== ACKNOWLEDGEMENTS =====
p = add_body(space_before=0, space_after=0)
r = p.add_run("Acknowledgements")
set_run(r, size=12, bold=True)
r = p.add_run(": The author would like to thank the supervising teacher, Lect. dr. ing. Stelian Sp\u00eenu, for guidance and support throughout this project.")
set_run(r, size=12)

# ===== KEYWORDS =====
p = add_body(space_before=0, space_after=0)
r = p.add_run("Keywords")
set_run(r, size=12, bold=True)
r = p.add_run(": ")
set_run(r, size=12)
r = p.add_run("visual impairment, assistive technology, depth estimation, object detection, face recognition, on-device AI")
set_run(r, size=12, italic=True)

# ===== REFERENCES =====
p = add_body(space_before=6, space_after=0)
r = p.add_run("References")
set_run(r, size=12, bold=True)
r = p.add_run(":")
set_run(r, size=12, bold=True)

refs = [
    '[1] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi. "You Only Look Once: Unified, Real-Time Object Detection." Proc. IEEE CVPR, Las Vegas, USA, Jun. 2016, pp. 779\u2013788.',
    '[2] L. Yang, B. Kang, Z. Huang, Z. Zhao, X. Xu, J. Feng, and H. Zhao. "Depth Anything V2." Advances in Neural Information Processing Systems, vol. 37, Dec. 2024.',
    '[3] W. Wu, H. Peng, and S. Yu. "YuNet: A Tiny Millisecond-level Face Detector." Machine Intelligence Research, vol. 20, no. 5, pp. 656\u2013665, 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(0.5)
    pf.first_line_indent = Cm(-0.5)
    r = p.add_run(ref)
    set_run(r, size=10)

output = "abstract_CERC2026.docx"
doc.save(output)
print(f"Saved: {output}")
